"""
Document processor for extracting and chunking text from various document formats
Optimized for accuracy with intelligent chunking strategies
"""

import aiohttp
import fitz  # PyMuPDF
import docx
import re
import nltk
from typing import List, Dict, Any, Tuple
from io import BytesIO
import logging

from src.core.config import settings

# Download required NLTK data with error handling
import ssl

try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context


def ensure_nltk_data():
    """Ensure all required NLTK data is available"""
    required_data = ["punkt", "punkt_tab", "stopwords"]

    for data_name in required_data:
        try:
            if data_name == "stopwords":
                nltk.data.find("corpora/stopwords")
            else:
                nltk.data.find(f"tokenizers/{data_name}")
        except LookupError:
            try:
                print(f"📥 Downloading NLTK {data_name}...")
                nltk.download(data_name, quiet=True)
                print(f"✅ {data_name} downloaded successfully")
            except Exception as e:
                print(f"❌ Failed to download {data_name}: {str(e)}")


# Ensure NLTK data is available
ensure_nltk_data()

logger = logging.getLogger(__name__)


class DocumentChunk:
    """Represents a chunk of document text with metadata"""

    def __init__(
        self,
        text: str,
        page_num: int = 0,
        chunk_id: str = "",
        metadata: Dict[str, Any] = None,
    ):
        self.text = text.strip()
        self.page_num = page_num
        self.chunk_id = chunk_id
        self.metadata = metadata or {}
        self.word_count = len(text.split())
        self.char_count = len(text)


class DocumentProcessor:
    """High-accuracy document processor with intelligent chunking"""

    def __init__(self):
        self.max_chunk_size = settings.MAX_CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.max_doc_size_mb = settings.MAX_DOCUMENT_SIZE_MB

    async def process_document(
        self, document_url: str
    ) -> Tuple[List[DocumentChunk], Dict[str, Any]]:
        """
        Process document from URL and return chunks with metadata

        Args:
            document_url: URL to the document

        Returns:
            Tuple of (chunks, metadata)
        """
        try:
            # Download document
            document_data, content_type = await self._download_document(document_url)

            # Extract text based on content type
            if "pdf" in content_type.lower():
                text, metadata = await self._extract_pdf_text(document_data)
            elif "word" in content_type.lower() or "docx" in content_type.lower():
                text, metadata = await self._extract_docx_text(document_data)
            else:
                # Try to extract as text
                text = document_data.decode("utf-8", errors="ignore")
                metadata = {"type": "text", "size": len(text)}

            # Add document URL to metadata
            metadata["source_url"] = document_url
            metadata["content_type"] = content_type

            # Create intelligent chunks
            chunks = await self._create_intelligent_chunks(text, metadata)

            logger.info(
                f"Processed document: {len(chunks)} chunks, {len(text)} characters"
            )

            return chunks, metadata

        except Exception as e:
            logger.error(f"Error processing document {document_url}: {str(e)}")
            raise

    async def _download_document(self, url: str) -> Tuple[bytes, str]:
        """Download document from URL"""
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(
                        f"Failed to download document: HTTP {response.status}"
                    )

                content_type = response.headers.get(
                    "content-type", "application/octet-stream"
                )
                data = await response.read()

                # Check file size
                if len(data) > self.max_doc_size_mb * 1024 * 1024:
                    raise Exception(f"Document too large: {len(data)} bytes")

                return data, content_type

    async def _extract_pdf_text(self, pdf_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyMuPDF"""
        doc = fitz.open(stream=pdf_data, filetype="pdf")

        text_parts = []
        metadata = {"type": "pdf", "pages": len(doc), "size": len(pdf_data)}

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()

            if page_text.strip():
                # Clean up text
                page_text = self._clean_text(page_text)
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

        doc.close()

        full_text = "\n\n".join(text_parts)
        metadata["extracted_chars"] = len(full_text)

        return full_text, metadata

    async def _extract_docx_text(self, docx_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        doc = docx.Document(BytesIO(docx_data))

        text_parts = []
        metadata = {
            "type": "docx",
            "paragraphs": len(doc.paragraphs),
            "size": len(docx_data),
        }

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        full_text = "\n\n".join(text_parts)
        metadata["extracted_chars"] = len(full_text)

        return full_text, metadata

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove special characters but keep punctuation
        text = re.sub(
            r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\~\`\|\\]",
            "",
            text,
        )

        # Fix common OCR errors (though we're told text is selectable)
        text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")

        return text.strip()

    async def _create_intelligent_chunks(
        self, text: str, metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Create intelligent chunks using semantic boundaries
        Prioritizes accuracy by preserving context
        """
        chunks = []

        # First, split by major sections (pages, chapters, etc.)
        sections = self._split_into_sections(text)

        chunk_id = 0
        for section_idx, section in enumerate(sections):
            if len(section.strip()) == 0:
                continue

            # For each section, create semantic chunks
            section_chunks = self._create_semantic_chunks(section, section_idx)

            for chunk_text in section_chunks:
                if len(chunk_text.strip()) > 50:  # Minimum chunk size
                    chunk = DocumentChunk(
                        text=chunk_text,
                        page_num=section_idx,
                        chunk_id=f"chunk_{chunk_id}",
                        metadata={
                            "section": section_idx,
                            "doc_type": metadata.get("type", "unknown"),
                        },
                    )
                    chunks.append(chunk)
                    chunk_id += 1

        return chunks

    def _split_into_sections(self, text: str) -> List[str]:
        """Split text into logical sections"""
        # Split by page markers first
        if "[Page " in text:
            sections = re.split(r"\[Page \d+\]", text)
            return [section.strip() for section in sections if section.strip()]

        # Split by double newlines (paragraph breaks)
        sections = text.split("\n\n")
        return [section.strip() for section in sections if section.strip()]

    def _create_semantic_chunks(self, text: str, section_idx: int) -> List[str]:
        """
        Create semantic chunks that preserve meaning
        Uses sentence boundaries and semantic coherence
        """
        if len(text) <= self.max_chunk_size:
            return [text]

        chunks = []

        # Split into sentences
        sentences = nltk.sent_tokenize(text)

        current_chunk = ""
        current_size = 0

        for sentence in sentences:
            sentence_size = len(sentence)

            # If adding this sentence would exceed max size
            if current_size + sentence_size > self.max_chunk_size and current_chunk:
                # Save current chunk
                chunks.append(current_chunk.strip())

                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_size = len(current_chunk)
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                current_size += sentence_size

        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        if len(text) <= self.chunk_overlap:
            return text

        # Get last few sentences for overlap
        sentences = nltk.sent_tokenize(text)
        overlap_text = ""

        for sentence in reversed(sentences):
            if len(overlap_text + sentence) <= self.chunk_overlap:
                overlap_text = sentence + " " + overlap_text
            else:
                break

        return overlap_text.strip()
