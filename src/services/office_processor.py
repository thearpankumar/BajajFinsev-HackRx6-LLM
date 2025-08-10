"""
Office Document Processor
Enhanced processing for DOCX/DOC and XLSX/XLS files with metadata extraction
"""

import logging
import time
from pathlib import Path
from typing import Any

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    import pandas as pd
    from openpyxl import load_workbook
    HAS_EXCEL_LIBS = True
except ImportError:
    HAS_EXCEL_LIBS = False

try:
    import xml.etree.ElementTree as ET
    import zipfile
    HAS_XML_SUPPORT = True
except ImportError:
    HAS_XML_SUPPORT = False

from src.core.config import config

logger = logging.getLogger(__name__)


class OfficeProcessor:
    """
    Enhanced Office document processor for DOCX/DOC and XLSX/XLS files
    Supports metadata extraction, table processing, and structured content analysis
    """

    def __init__(self):
        self.max_document_size_mb = config.max_document_size_mb
        # Use proper config access with defaults instead of getattr
        self.excel_max_rows = config.excel_max_rows if hasattr(config, 'excel_max_rows') else 10000
        self.excel_sheet_limit = config.excel_sheet_limit if hasattr(config, 'excel_sheet_limit') else 5

        # Processing statistics
        self.total_processed = 0
        self.total_processing_time = 0.0

        self._check_dependencies()
        logger.info("OfficeProcessor initialized for DOCX/XLSX processing")

    def _check_dependencies(self):
        """Check available dependencies"""
        deps = {
            "python-docx": HAS_DOCX,
            "pandas": HAS_EXCEL_LIBS,
            "xml_support": HAS_XML_SUPPORT
        }

        logger.info("📦 Office processing dependencies:")
        for dep, available in deps.items():
            status = "✅" if available else "❌"
            logger.info(f"  {status} {dep}")

    async def process_document(self, file_path: str, file_type: str) -> dict[str, Any]:
        """
        Process office document based on type
        
        Args:
            file_path: Path to document file
            file_type: Document type (docx, doc, xlsx, xls, csv)
            
        Returns:
            Processing results with extracted content and metadata
        """
        filepath = Path(file_path)

        if not filepath.exists():
            return {
                "status": "error",
                "error": f"File not found: {file_path}",
                "file_path": file_path
            }

        # Check file size
        file_size_mb = filepath.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_document_size_mb:
            return {
                "status": "error",
                "error": f"File too large: {file_size_mb:.1f}MB > {self.max_document_size_mb}MB",
                "file_path": file_path
            }

        logger.info(f"🔄 Processing {file_type.upper()} document: {filepath.name}")

        start_time = time.time()

        try:
            # Route to appropriate processor
            if file_type in ['docx', 'doc']:
                result = await self._process_word_document(filepath, file_type)
            elif file_type in ['xlsx', 'xls', 'csv']:
                result = await self._process_excel_document(filepath, file_type)
            else:
                return {
                    "status": "error",
                    "error": f"Unsupported file type: {file_type}",
                    "file_path": file_path
                }

            # Add common metadata
            processing_time = time.time() - start_time
            self.total_processed += 1
            self.total_processing_time += processing_time

            if result["status"] == "success":
                result.update({
                    "file_path": file_path,
                    "file_size_mb": round(file_size_mb, 2),
                    "processing_time": round(processing_time, 2),
                    "processor": "OfficeProcessor"
                })

            logger.info(f"✅ {file_type.upper()} processed in {processing_time:.2f}s")
            return result

        except Exception as e:
            error_msg = f"{file_type.upper()} processing failed: {str(e)}"
            logger.error(f"❌ {error_msg}")
            return {
                "status": "error",
                "error": error_msg,
                "file_path": file_path,
                "file_type": file_type
            }

    async def _process_word_document(self, filepath: Path, file_type: str) -> dict[str, Any]:
        """Process DOCX document with advanced features"""
        if not HAS_DOCX:
            return {
                "status": "error",
                "error": "python-docx not available. Install with: pip install python-docx"
            }

        try:
            # Load document
            doc = DocxDocument(str(filepath))

            # Extract metadata
            metadata = self._extract_word_metadata(doc)

            # Extract content structure
            content_structure = await self._extract_word_structure(doc)

            # Extract tables
            tables = self._extract_word_tables(doc)

            # Combine full text
            full_text = self._combine_word_content(content_structure, tables)

            return {
                "status": "success",
                "document_type": "word_document",
                "metadata": metadata,
                "content": {
                    "full_text": full_text,
                    "structure": content_structure,
                    "tables": tables,
                    "char_count": len(full_text),
                    "word_count": len(full_text.split()),
                    "paragraph_count": len([p for p in content_structure if p["type"] == "paragraph"]),
                    "table_count": len(tables)
                },
                "extraction_features": {
                    "text_extraction": True,
                    "table_extraction": True,
                    "metadata_extraction": True,
                    "structure_analysis": True,
                    "style_analysis": True
                }
            }

        except Exception as e:
            return {
                "status": "error",
                "error": f"DOCX processing error: {str(e)}"
            }

    def _extract_word_metadata(self, doc: 'DocxDocument') -> dict[str, Any]:
        """Extract comprehensive metadata from Word document"""
        try:
            core_props = doc.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "category": core_props.category or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "version": core_props.version or "",
                "language": core_props.language or ""
            }

            # Add document statistics
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "has_header_footer": any(
                    section.header.paragraphs or section.footer.paragraphs
                    for section in doc.sections
                )
            })

            return metadata

        except Exception as e:
            logger.warning(f"Word metadata extraction failed: {str(e)}")
            return {"extraction_error": str(e)}

    async def _extract_word_structure(self, doc: 'DocxDocument') -> list[dict[str, Any]]:
        """Extract structured content from Word document"""
        structure = []

        for i, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()

            if not para_text:  # Skip empty paragraphs
                continue

            # Analyze paragraph style and formatting
            style_info = self._analyze_paragraph_style(paragraph)

            para_info = {
                "type": "paragraph",
                "index": i,
                "text": para_text,
                "char_count": len(para_text),
                "word_count": len(para_text.split()),
                "style": style_info,
                "is_heading": style_info.get("is_heading", False),
                "heading_level": style_info.get("heading_level", 0)
            }

            structure.append(para_info)

        return structure

    def _analyze_paragraph_style(self, paragraph) -> dict[str, Any]:
        """Analyze paragraph style and formatting"""
        try:
            style_name = paragraph.style.name if paragraph.style else "Normal"

            # Detect headings
            is_heading = "Heading" in style_name or "Title" in style_name
            heading_level = 0

            if is_heading and "Heading" in style_name:
                try:
                    heading_level = int(style_name.split()[-1])
                except:
                    heading_level = 1

            # Analyze runs for formatting
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            has_italic = any(run.italic for run in paragraph.runs if run.italic)
            has_underline = any(run.underline for run in paragraph.runs if run.underline)

            return {
                "style_name": style_name,
                "is_heading": is_heading,
                "heading_level": heading_level,
                "has_bold": has_bold,
                "has_italic": has_italic,
                "has_underline": has_underline,
                "alignment": str(paragraph.alignment) if paragraph.alignment else "unknown"
            }

        except Exception as e:
            return {
                "style_name": "unknown",
                "analysis_error": str(e)
            }

    def _extract_word_tables(self, doc: 'DocxDocument') -> list[dict[str, Any]]:
        """Extract tables from Word document"""
        tables = []

        for i, table in enumerate(doc.tables):
            try:
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if not table_data:
                    continue

                # Convert to text format
                table_text = self._table_data_to_text(table_data, f"Table_{i+1}")

                table_info = {
                    "table_id": f"word_table_{i+1}",
                    "table_index": i,
                    "row_count": len(table_data),
                    "col_count": len(table_data[0]) if table_data else 0,
                    "raw_data": table_data,
                    "formatted_text": table_text,
                    "extraction_method": "python_docx"
                }

                tables.append(table_info)

            except Exception as e:
                logger.warning(f"Word table extraction failed for table {i+1}: {str(e)}")
                continue

        return tables

    def _combine_word_content(self, structure: list[dict], tables: list[dict]) -> str:
        """Combine Word document content"""
        content_parts = []

        # Add structured paragraphs
        for item in structure:
            if item["type"] == "paragraph":
                # Add heading markers
                if item["is_heading"] and item["heading_level"] > 0:
                    level_marker = "#" * item["heading_level"]
                    content_parts.append(f"{level_marker} {item['text']}")
                else:
                    content_parts.append(item["text"])

        # Add tables
        for table in tables:
            content_parts.append(f"\n{table['formatted_text']}")

        return "\n\n".join(content_parts)

    async def _process_excel_document(self, filepath: Path, file_type: str) -> dict[str, Any]:
        """Process Excel document with enhanced features"""
        if not HAS_EXCEL_LIBS:
            return {
                "status": "error",
                "error": "Excel libraries not available. Install with: pip install pandas openpyxl xlrd"
            }

        try:
            if file_type == 'csv':
                return await self._process_csv_file(filepath)
            else:
                return await self._process_excel_file(filepath, file_type)

        except Exception as e:
            return {
                "status": "error",
                "error": f"Excel processing error: {str(e)}"
            }

    async def _process_csv_file(self, filepath: Path) -> dict[str, Any]:
        """Process CSV file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            df = None

            for encoding in encodings:
                try:
                    df = pd.read_csv(str(filepath), encoding=encoding, nrows=self.excel_max_rows)
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                return {
                    "status": "error",
                    "error": "Could not read CSV file with any supported encoding"
                }

            # Process the DataFrame
            sheet_text = self._dataframe_to_text(df, "CSV_Data")

            full_text = f"=== CSV FILE ===\n{sheet_text}"

            return {
                "status": "success",
                "document_type": "csv_file",
                "metadata": {
                    "file_type": "csv",
                    "row_count": len(df),
                    "column_count": len(df.columns),
                    "columns": list(df.columns),
                    "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
                },
                "content": {
                    "full_text": full_text,
                    "sheet_count": 1,
                    "sheets": [{
                        "sheet_name": "CSV_Data",
                        "text": sheet_text,
                        "row_count": len(df),
                        "col_count": len(df.columns),
                        "columns": list(df.columns)
                    }],
                    "char_count": len(full_text),
                    "word_count": len(full_text.split())
                },
                "extraction_features": {
                    "data_extraction": True,
                    "type_inference": True,
                    "metadata_extraction": True
                }
            }

        except Exception as e:
            return {
                "status": "error",
                "error": f"CSV processing failed: {str(e)}"
            }

    async def _process_excel_file(self, filepath: Path, file_type: str) -> dict[str, Any]:
        """Process Excel file (XLSX/XLS)"""
        try:
            # Load workbook for metadata
            if file_type == 'xlsx':
                workbook = load_workbook(str(filepath), read_only=True, data_only=True)
                metadata = self._extract_excel_metadata(workbook, filepath)
            else:
                metadata = {"file_type": file_type, "extraction_method": "pandas"}

            # Load with pandas for data processing
            excel_file = pd.ExcelFile(str(filepath))

            # Limit number of sheets processed
            sheet_names = excel_file.sheet_names[:self.excel_sheet_limit]

            sheets_data = []
            all_text_parts = []

            for sheet_name in sheet_names:
                try:
                    # Load sheet data
                    df = pd.read_excel(excel_file, sheet_name=sheet_name, nrows=self.excel_max_rows)

                    if df.empty:
                        continue

                    # Convert to text
                    sheet_text = self._dataframe_to_text(df, sheet_name)

                    sheet_info = {
                        "sheet_name": sheet_name,
                        "text": sheet_text,
                        "row_count": len(df),
                        "col_count": len(df.columns),
                        "columns": list(df.columns),
                        "data_types": {col: str(dtype) for col, dtype in df.dtypes.items()},
                        "has_data": not df.empty
                    }

                    sheets_data.append(sheet_info)
                    all_text_parts.append(f"=== SHEET: {sheet_name} ===\n{sheet_text}")

                except Exception as e:
                    logger.warning(f"Failed to process sheet '{sheet_name}': {str(e)}")
                    continue

            full_text = "\n\n".join(all_text_parts)

            return {
                "status": "success",
                "document_type": "excel_file",
                "metadata": metadata,
                "content": {
                    "full_text": full_text,
                    "sheet_count": len(sheets_data),
                    "sheets": sheets_data,
                    "char_count": len(full_text),
                    "word_count": len(full_text.split()),
                    "total_rows": sum(sheet["row_count"] for sheet in sheets_data),
                    "total_columns": sum(sheet["col_count"] for sheet in sheets_data)
                },
                "extraction_features": {
                    "multi_sheet_processing": True,
                    "data_type_inference": True,
                    "metadata_extraction": file_type == 'xlsx',
                    "row_limit_applied": self.excel_max_rows,
                    "sheet_limit_applied": self.excel_sheet_limit
                }
            }

        except Exception as e:
            return {
                "status": "error",
                "error": f"Excel file processing failed: {str(e)}"
            }

    def _extract_excel_metadata(self, workbook, filepath: Path) -> dict[str, Any]:
        """Extract metadata from Excel workbook"""
        try:
            props = workbook.properties

            metadata = {
                "title": props.title or "",
                "creator": props.creator or "",
                "subject": props.subject or "",
                "description": props.description or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.lastModifiedBy or "",
                "version": props.version or "",
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "file_type": "xlsx"
            }

            return metadata

        except Exception as e:
            logger.warning(f"Excel metadata extraction failed: {str(e)}")
            return {
                "file_type": "xlsx",
                "sheet_names": workbook.sheetnames if hasattr(workbook, 'sheetnames') else [],
                "extraction_error": str(e)
            }

    def _detect_column_types(self, df: 'pd.DataFrame') -> dict[int, str]:
        """
        Dynamically detect column types based on content patterns
        Returns a mapping of column_index -> meaningful_name
        """
        detected_columns = {}
        
        # Sample some rows to analyze patterns
        sample_size = min(10, len(df))
        sample_df = df.head(sample_size)
        
        for col_idx, col_name in enumerate(df.columns):
            # Get sample values for this column (non-null, non-empty)
            sample_values = []
            for val in sample_df.iloc[:, col_idx].dropna():
                if str(val).strip() and str(val).strip() != '':
                    sample_values.append(str(val).strip())
            
            if not sample_values:
                detected_columns[col_idx] = f"Column{col_idx+1}"
                continue
            
            # Detect patterns in the sample values
            column_type = self._classify_column_content(sample_values, col_name)
            detected_columns[col_idx] = column_type
        
        return detected_columns
    
    def _classify_column_content(self, sample_values: list[str], original_name: str) -> str:
        """Classify column content based on patterns"""
        import re
        
        # Clean original name first
        clean_original = self._clean_column_name(original_name)
        
        # If original name is already meaningful, use it
        meaningful_names = [
            'name', 'firstname', 'lastname', 'fullname', 'person', 'customer',
            'phone', 'mobile', 'telephone', 'contact', 'number', 
            'pincode', 'zip', 'postal', 'code', 'pin',
            'salary', 'amount', 'price', 'cost', 'value', 'income',
            'email', 'address', 'location', 'city', 'state', 'country',
            'id', 'identifier', 'reference', 'ref',
            'date', 'time', 'timestamp', 'created', 'updated',
            'status', 'type', 'category', 'group'
        ]
        
        clean_lower = clean_original.lower()
        for meaningful in meaningful_names:
            if meaningful in clean_lower:
                return clean_original
        
        # Pattern-based detection
        numeric_count = 0
        text_count = 0
        phone_count = 0
        pincode_count = 0
        name_count = 0
        
        for value in sample_values:
            # Check for phone numbers (10+ digits, may have spaces/dashes)
            if re.match(r'^[\d\s\-\+\(\)]{10,}$', value):
                phone_count += 1
            
            # Check for pincode/zip patterns (4-8 digits)
            elif re.match(r'^\d{4,8}$', value):
                pincode_count += 1
            
            # Check for pure numbers (salary, amounts, etc.)
            elif re.match(r'^\d+(\.\d+)?$', value):
                numeric_count += 1
            
            # Check for names (2+ words with letters)
            elif re.match(r'^[a-zA-Z\s]{2,}$', value) and len(value.split()) >= 2:
                name_count += 1
            
            # Everything else is text
            else:
                text_count += 1
        
        total_samples = len(sample_values)
        if total_samples == 0:
            return "Column"
        
        # Determine column type based on majority pattern
        phone_ratio = phone_count / total_samples
        pincode_ratio = pincode_count / total_samples
        name_ratio = name_count / total_samples
        numeric_ratio = numeric_count / total_samples
        
        # Use thresholds to classify (at least 70% of samples match pattern)
        if phone_ratio >= 0.7:
            return "Mobile_Number"
        elif pincode_ratio >= 0.7:
            return "Pincode"
        elif name_ratio >= 0.7:
            return "Name"
        elif numeric_ratio >= 0.7:
            return "Amount"
        else:
            # If no clear pattern, use cleaned original name
            return clean_original if clean_original != "Column" else "Text_Data"
    
    def _clean_column_name(self, col_name: str) -> str:
        """Clean and normalize column names"""
        import re
        
        clean_name = str(col_name).strip()
        
        # Handle common problematic patterns
        if clean_name.startswith('Unnamed:'):
            return "Column"
        
        if clean_name.startswith('From: System Administrator'):
            return "Column"
        
        if clean_name in ['', ' ', 'nan', 'NaN', 'None']:
            return "Column"
        
        # Clean up the name - remove special characters, normalize spaces
        clean_name = re.sub(r'[^\w\s]', ' ', clean_name)
        clean_name = re.sub(r'\s+', '_', clean_name.strip())
        
        # Capitalize first letter of each word
        clean_name = '_'.join(word.capitalize() for word in clean_name.split('_') if word)
        
        return clean_name if clean_name else "Column"

    def _detect_table_column_types(self, headers: list[str], data_rows: list[list[str]]) -> dict[int, str]:
        """Detect column types for table data"""
        detected_columns = {}
        
        # Sample some rows for analysis
        sample_size = min(10, len(data_rows))
        sample_rows = data_rows[:sample_size]
        
        for col_idx, header in enumerate(headers):
            # Get sample values for this column
            sample_values = []
            for row in sample_rows:
                if col_idx < len(row) and str(row[col_idx]).strip():
                    sample_values.append(str(row[col_idx]).strip())
            
            if not sample_values:
                detected_columns[col_idx] = f"Column{col_idx+1}"
                continue
            
            # Use the same classification logic
            column_type = self._classify_column_content(sample_values, header)
            detected_columns[col_idx] = column_type
        
        return detected_columns

    def _find_actual_headers(self, df: 'pd.DataFrame') -> list[str]:
        """Find actual column headers in the data by looking for header-like rows"""
        # Look through more rows to find proper headers, but be more selective
        header_candidates = []
        
        for idx in range(min(20, len(df))):  # Search more rows
            row = df.iloc[idx]
            row_values = [str(val).strip() for val in row if str(val).strip() and str(val).strip().lower() != 'nan']
            
            # Skip empty rows
            if not row_values:
                continue
            
            # Skip rows that are clearly not headers (too long text, sentences)
            if any(len(str(val).strip()) > 50 for val in row_values):
                continue
            
            # Check if this row looks like headers with exact keyword matching
            header_score = 0
            exact_header_words = ['name', 'mobile', 'phone', 'pincode', 'salary', 'amount', 'id', 'email', 'address']
            
            for val in row_values:
                val_lower = str(val).strip().lower()
                # Exact word matching or very close matches
                if val_lower in exact_header_words:
                    header_score += 2  # Higher score for exact matches
                elif any(hw in val_lower and len(val_lower) <= 20 for hw in exact_header_words):
                    header_score += 1  # Lower score for partial matches in short strings
            
            # Require a strong match - at least 3 columns that look like headers
            if header_score >= 3 and len(row_values) >= 3:
                # This row contains headers
                clean_headers = []
                for col_idx, val in enumerate(row):
                    header_val = str(val).strip()
                    if header_val and header_val.lower() != 'nan':
                        clean_headers.append(self._clean_header_name(header_val))
                    else:
                        # Fallback for empty headers
                        clean_headers.append(f"Column{col_idx+1}")
                
                if len(clean_headers) > 0:
                    header_candidates.append((header_score, clean_headers, idx))
        
        # Return the best header candidate (highest score, preferring later rows if tied)
        if header_candidates:
            header_candidates.sort(key=lambda x: (x[0], x[2]), reverse=True)
            return header_candidates[0][1]
        
        return None

    def _clean_header_name(self, header: str) -> str:
        """Clean header names for better readability"""
        header = str(header).strip()
        
        if not header or header.lower() in ['nan', 'none', '']:
            return 'Column'
        
        header = header.replace('_', ' ').replace('-', ' ')
        
        # Map common variations - be more specific
        header_lower = header.lower()
        if header_lower == 'name' or 'person' in header_lower:
            return 'Name'
        elif 'mobile' in header_lower or 'phone' in header_lower:
            return 'Mobile_Number'
        elif 'pincode' in header_lower or 'pin code' in header_lower:
            return 'Pincode'  # This is the key mapping!
        elif header_lower == 'salary' or 'amount' in header_lower or 'income' in header_lower:
            return 'Salary'
        elif 'email' in header_lower:
            return 'Email'
        elif 'address' in header_lower:
            return 'Address'
        elif 'id' in header_lower:
            return 'ID'
        else:
            # Keep original format but clean it up
            clean = '_'.join(word.capitalize() for word in header.split() if word)
            return clean if clean else 'Column'

    def _detect_column_type_by_content(self, sample_values, original_name: str) -> str:
        """Detect column type by analyzing content patterns"""
        import re
        
        sample_list = [str(val).strip() for val in sample_values if str(val).strip()]
        if not sample_list:
            return "Column"
        
        # Pattern matching
        phone_count = sum(1 for val in sample_list if re.match(r'^\d{10,}$', val))
        pincode_count = sum(1 for val in sample_list if re.match(r'^\d{4,8}$', val))
        name_count = sum(1 for val in sample_list if re.match(r'^[a-zA-Z\s]{2,}$', val) and len(val.split()) >= 2)
        numeric_count = sum(1 for val in sample_list if re.match(r'^\d+(\.\d+)?$', val))
        
        total = len(sample_list)
        if total == 0:
            return "Column"
        
        # Determine type (70% threshold)
        if phone_count / total >= 0.7:
            return "Mobile_Number"
        elif pincode_count / total >= 0.7:
            return "Pincode"
        elif name_count / total >= 0.7:
            return "Name"
        elif numeric_count / total >= 0.7:
            return "Amount"
        else:
            return self._clean_column_name(original_name)

    def _dataframe_to_text(self, df: 'pd.DataFrame', sheet_name: str) -> str:
        """Convert DataFrame to readable text"""
        try:
            # Clean DataFrame
            df = df.fillna("")

            # Create header
            headers = " | ".join(str(col) for col in df.columns)
            text_parts = [
                f"Sheet: {sheet_name}",
                f"Columns ({len(df.columns)}): {headers}",
                f"Rows: {len(df)}",
                ""
            ]

            # Look for actual column headers in the data
            actual_headers = self._find_actual_headers(df)
            
            # If no headers found, try to detect them from a specific row that looks like headers
            if not actual_headers:
                for idx in range(min(15, len(df))):
                    row = df.iloc[idx]
                    row_values = [str(val).strip().lower() for val in row if str(val).strip() and str(val).strip() != 'nan']
                    
                    # Check if this row has header-like content
                    header_keywords = sum(1 for val in row_values if any(kw in val for kw in ['name', 'mobile', 'phone', 'pincode', 'salary']))
                    
                    if header_keywords >= 3:  # At least 3 header-like words
                        # Use this row as headers
                        actual_headers = []
                        for col_idx, val in enumerate(row):
                            header_val = str(val).strip()
                            if header_val and header_val.lower() != 'nan':
                                actual_headers.append(self._clean_header_name(header_val))
                            else:
                                actual_headers.append(f"Column{col_idx+1}")
                        break
            
            # Add all data rows with proper column naming
            for idx, row in df.iterrows():
                # Create row with column_name=value format for better AI understanding
                row_parts = []
                for col_idx, (col, val) in enumerate(row.items()):
                    # Use actual header names if found, otherwise detect dynamically
                    if actual_headers and col_idx < len(actual_headers):
                        clean_col = actual_headers[col_idx]
                    else:
                        clean_col = self._detect_column_type_by_content(df.iloc[:, col_idx].dropna().head(10), str(col))
                    
                    row_parts.append(f"{clean_col}={str(val).strip()}")
                
                row_text = " | ".join(row_parts)
                text_parts.append(f"Row {idx + 1}: {row_text}")

            return "\n".join(text_parts)

        except Exception as e:
            logger.warning(f"DataFrame to text conversion failed: {str(e)}")
            return f"Sheet: {sheet_name}\nConversion failed: {str(e)}"

    def _table_data_to_text(self, table_data: list[list[str]], table_name: str) -> str:
        """Convert table data to text format"""
        try:
            if not table_data:
                return f"Table: {table_name}\nNo data"

            # Assume first row is headers
            headers = table_data[0]
            data_rows = table_data[1:]

            text_parts = [
                f"=== {table_name} ===",
                f"Columns: {' | '.join(headers)}",
                ""
            ]

            # Detect column types dynamically from table data
            detected_table_columns = self._detect_table_column_types(headers, data_rows)
            
            # Add data rows with intelligent column naming
            for i, row in enumerate(data_rows):
                # Create row with column_name=value format for better AI understanding
                row_parts = []
                for j, cell in enumerate(row):
                    col_name = detected_table_columns.get(j, f"Column{j+1}")
                    row_parts.append(f"{col_name}={str(cell).strip()}")
                
                row_text = " | ".join(row_parts)
                text_parts.append(f"Row {i + 1}: {row_text}")

            return "\n".join(text_parts)

        except Exception as e:
            return f"Table: {table_name}\nProcessing failed: {str(e)}"

    def get_processing_stats(self) -> dict[str, Any]:
        """Get processing statistics"""
        avg_processing_time = (
            self.total_processing_time / self.total_processed
            if self.total_processed > 0 else 0.0
        )

        return {
            "total_documents_processed": self.total_processed,
            "total_processing_time": round(self.total_processing_time, 2),
            "average_processing_time": round(avg_processing_time, 2),
            "supported_formats": ["docx", "doc", "xlsx", "xls", "csv"],
            "features_available": {
                "word_processing": HAS_DOCX,
                "excel_processing": HAS_EXCEL_LIBS,
                "metadata_extraction": True,
                "table_extraction": True,
                "structure_analysis": HAS_DOCX
            },
            "processing_limits": {
                "max_excel_rows": self.excel_max_rows,
                "max_excel_sheets": self.excel_sheet_limit,
                "max_file_size_mb": self.max_document_size_mb
            }
        }
