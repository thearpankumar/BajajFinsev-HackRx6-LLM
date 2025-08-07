"""
Text Extraction Service for BajajFinsev Hybrid RAG System
Extracts text from multiple document formats: PDF, DOCX, XLSX, Images
Supports OCR, structured data extraction, and text preprocessing
"""

import asyncio
import logging
import io
import os
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass
from pathlib import Path
import time
import tempfile

# Document processing libraries
import fitz  # PyMuPDF for PDF processing
from docx import Document  # python-docx for DOCX
import pandas as pd  # For Excel files
import csv
from PIL import Image
import pytesseract  # OCR
import cv2
import numpy as np

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of text extraction operation"""
    success: bool
    text: str = ""
    metadata: Dict[str, Any] = None
    page_count: int = 0
    word_count: int = 0
    extraction_time: float = 0.0
    file_format: str = ""
    error_message: Optional[str] = None
    structured_data: Optional[Dict] = None


class TextExtractor:
    """
    Universal text extraction service supporting multiple formats
    Features: PDF, DOCX, XLSX, CSV, Images with OCR, text preprocessing
    """
    
    def __init__(self, 
                 enable_ocr: bool = True,
                 ocr_language: str = 'eng',
                 max_image_size: int = 10 * 1024 * 1024,  # 10MB
                 preprocessing_enabled: bool = True):
        """
        Initialize text extractor
        
        Args:
            enable_ocr: Whether to enable OCR for images and scanned PDFs
            ocr_language: Tesseract language code
            max_image_size: Maximum image size for OCR processing
            preprocessing_enabled: Whether to enable text preprocessing
        """
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language
        self.max_image_size = max_image_size
        self.preprocessing_enabled = preprocessing_enabled
        
        # Extraction statistics
        self.stats = {
            'total_extractions': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'avg_extraction_time': 0.0,
            'formats_processed': {},
            'total_pages_processed': 0,
            'total_words_extracted': 0,
            'ocr_operations': 0
        }
        
        # Initialize OCR if enabled
        if self.enable_ocr:
            try:
                # Test OCR availability
                pytesseract.get_tesseract_version()
                logger.info(f"✅ OCR enabled with language: {ocr_language}")
            except Exception as e:
                logger.warning(f"⚠️ OCR initialization failed: {e}")
                self.enable_ocr = False
        
        logger.info("TextExtractor initialized")
        logger.info(f"OCR enabled: {self.enable_ocr}")
        logger.info(f"Preprocessing enabled: {self.preprocessing_enabled}")
    
    async def extract_text(self, file_path: str) -> ExtractionResult:
        """
        Extract text from document based on file extension
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        start_time = time.time()
        self.stats['total_extractions'] += 1
        
        logger.info(f"📄 Starting text extraction: {file_path}")
        
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                return self._create_error_result(
                    f"File not found: {file_path}",
                    time.time() - start_time
                )
            
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            logger.debug(f"File extension: {file_ext}, Size: {file_size:,} bytes")
            
            # Route to appropriate extractor
            if file_ext == '.pdf':
                result = await self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                result = await self._extract_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                result = await self._extract_from_excel(file_path)
            elif file_ext == '.csv':
                result = await self._extract_from_csv(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
                result = await self._extract_from_image(file_path)
            elif file_ext == '.txt':
                result = await self._extract_from_text(file_path)
            else:
                return self._create_error_result(
                    f"Unsupported file format: {file_ext}",
                    time.time() - start_time
                )
            
            # Update extraction time
            result.extraction_time = time.time() - start_time
            result.file_format = file_ext
            
            # Update statistics
            if result.success:
                self.stats['successful_extractions'] += 1
                self.stats['total_pages_processed'] += result.page_count
                self.stats['total_words_extracted'] += result.word_count
                self._update_format_stats(file_ext)
                self._update_avg_extraction_time(result.extraction_time)
                
                logger.info(f"✅ Extraction successful: {result.word_count} words, "
                           f"{result.page_count} pages in {result.extraction_time:.2f}s")
            else:
                self.stats['failed_extractions'] += 1
                logger.error(f"❌ Extraction failed: {result.error_message}")
            
            return result
            
        except Exception as e:
            error_msg = f"Unexpected error during extraction: {str(e)}"
            logger.error(f"❌ {error_msg}")
            self.stats['failed_extractions'] += 1
            
            return self._create_error_result(error_msg, time.time() - start_time)
    
    async def _extract_from_pdf(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF with OCR fallback"""
        try:
            # Open PDF document
            doc = fitz.open(file_path)
            
            full_text = []
            page_count = len(doc)
            metadata = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "page_count": page_count
            }
            
            logger.debug(f"PDF opened: {page_count} pages")
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                
                # Try text extraction first
                text = page.get_text()
                
                # If text is minimal and OCR is enabled, try OCR
                if len(text.strip()) < 50 and self.enable_ocr:
                    logger.debug(f"Page {page_num + 1}: Minimal text found, trying OCR")
                    
                    try:
                        # Render page as image
                        mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Perform OCR
                        ocr_text = await self._perform_ocr(image)
                        if len(ocr_text.strip()) > len(text.strip()):
                            text = ocr_text
                            self.stats['ocr_operations'] += 1
                            logger.debug(f"Page {page_num + 1}: OCR improved text extraction")
                    
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page {page_num + 1}: {ocr_error}")
                
                if text.strip():
                    full_text.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
            
            doc.close()
            
            # Combine text
            combined_text = "\n\n".join(full_text)
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                combined_text = self._preprocess_text(combined_text)
            
            return ExtractionResult(
                success=True,
                text=combined_text,
                metadata=metadata,
                page_count=page_count,
                word_count=len(combined_text.split()),
                structured_data={"pages": full_text}
            )
            
        except Exception as e:
            error_msg = f"PDF extraction failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _extract_from_docx(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX using python-docx"""
        try:
            # Open DOCX document
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Only add non-empty rows
                        table_data.append(" | ".join(row_data))
                
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine text
            full_text_parts = []
            
            if paragraphs:
                full_text_parts.append("=== Document Content ===")
                full_text_parts.extend(paragraphs)
            
            if tables_text:
                full_text_parts.append("\n=== Tables ===")
                full_text_parts.extend(tables_text)
            
            combined_text = "\n\n".join(full_text_parts)
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                combined_text = self._preprocess_text(combined_text)
            
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_text),
                "page_count": 1  # DOCX doesn't have clear page boundaries
            }
            
            return ExtractionResult(
                success=True,
                text=combined_text,
                metadata=metadata,
                page_count=1,
                word_count=len(combined_text.split()),
                structured_data={"paragraphs": paragraphs, "tables": tables_text}
            )
            
        except Exception as e:
            error_msg = f"DOCX extraction failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _extract_from_excel(self, file_path: str) -> ExtractionResult:
        """Extract text from Excel files using pandas"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            sheets_text = []
            total_rows = 0
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                # Convert DataFrame to text
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                
                # Add column headers
                headers = " | ".join(str(col) for col in df.columns)
                sheet_text += f"Headers: {headers}\n\n"
                
                # Add data rows
                for idx, row in df.iterrows():
                    row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                    if row_text.strip():
                        sheet_text += f"Row {idx + 1}: {row_text}\n"
                        total_rows += 1
                
                sheets_text.append(sheet_text)
            
            # Combine all sheets
            combined_text = "\n\n".join(sheets_text)
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                combined_text = self._preprocess_text(combined_text)
            
            metadata = {
                "sheet_count": len(excel_data),
                "total_rows": total_rows,
                "page_count": len(excel_data)  # Each sheet as a page
            }
            
            return ExtractionResult(
                success=True,
                text=combined_text,
                metadata=metadata,
                page_count=len(excel_data),
                word_count=len(combined_text.split()),
                structured_data={"sheets": list(excel_data.keys())}
            )
            
        except Exception as e:
            error_msg = f"Excel extraction failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _extract_from_csv(self, file_path: str) -> ExtractionResult:
        """Extract text from CSV files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Could not decode CSV file with any encoding")
            
            # Convert to text
            text_parts = []
            
            # Add headers
            headers = " | ".join(str(col) for col in df.columns)
            text_parts.append(f"CSV Headers: {headers}\n")
            
            # Add data rows
            for idx, row in df.iterrows():
                row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                if row_text.strip():
                    text_parts.append(f"Row {idx + 1}: {row_text}")
            
            combined_text = "\n".join(text_parts)
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                combined_text = self._preprocess_text(combined_text)
            
            metadata = {
                "row_count": len(df),
                "column_count": len(df.columns),
                "page_count": 1
            }
            
            return ExtractionResult(
                success=True,
                text=combined_text,
                metadata=metadata,
                page_count=1,
                word_count=len(combined_text.split()),
                structured_data={"columns": list(df.columns)}
            )
            
        except Exception as e:
            error_msg = f"CSV extraction failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _extract_from_image(self, file_path: str) -> ExtractionResult:
        """Extract text from images using OCR"""
        if not self.enable_ocr:
            return self._create_error_result("OCR is disabled", 0)
        
        try:
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_image_size:
                return self._create_error_result(
                    f"Image too large: {file_size} bytes > {self.max_image_size} bytes", 0
                )
            
            # Load image
            image = Image.open(file_path)
            
            # Perform OCR
            text = await self._perform_ocr(image)
            self.stats['ocr_operations'] += 1
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                text = self._preprocess_text(text)
            
            metadata = {
                "image_size": image.size,
                "image_mode": image.mode,
                "page_count": 1
            }
            
            return ExtractionResult(
                success=True,
                text=text,
                metadata=metadata,
                page_count=1,
                word_count=len(text.split())
            )
            
        except Exception as e:
            error_msg = f"Image OCR failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _extract_from_text(self, file_path: str) -> ExtractionResult:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with any encoding")
            
            # Preprocess if enabled
            if self.preprocessing_enabled:
                text = self._preprocess_text(text)
            
            # Count lines as pages
            line_count = len(text.split('\n'))
            page_count = max(1, line_count // 50)  # Approximate 50 lines per page
            
            metadata = {
                "line_count": line_count,
                "character_count": len(text),
                "page_count": page_count
            }
            
            return ExtractionResult(
                success=True,
                text=text,
                metadata=metadata,
                page_count=page_count,
                word_count=len(text.split())
            )
            
        except Exception as e:
            error_msg = f"Text file extraction failed: {str(e)}"
            return self._create_error_result(error_msg, 0)
    
    async def _perform_ocr(self, image: Image.Image) -> str:
        """Perform OCR on PIL Image with preprocessing"""
        try:
            # Convert PIL Image to OpenCV format for preprocessing
            cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Image preprocessing for better OCR
            gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
            
            # Noise removal
            denoised = cv2.medianBlur(gray, 5)
            
            # Thresholding
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Convert back to PIL Image
            processed_image = Image.fromarray(thresh)
            
            # Perform OCR
            ocr_config = f'--oem 3 --psm 6 -l {self.ocr_language}'
            text = pytesseract.image_to_string(processed_image, config=ocr_config)
            
            return text.strip()
            
        except Exception as e:
            logger.warning(f"OCR preprocessing failed, trying direct OCR: {e}")
            
            # Fallback to direct OCR
            try:
                text = pytesseract.image_to_string(image, lang=self.ocr_language)
                return text.strip()
            except Exception as e2:
                logger.error(f"Direct OCR also failed: {e2}")
                return ""
    
    def _preprocess_text(self, text: str) -> str:
        """Preprocess extracted text for better quality"""
        if not text:
            return text
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove multiple newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        # Remove page break artifacts
        text = text.replace('\f', '\n')  # Form feed
        text = text.replace('\r', '')    # Carriage return
        
        # Clean up common OCR artifacts
        text = text.replace('|', 'I')  # Common OCR confusion
        text = text.replace('0', 'O')  # In some contexts
        
        return text.strip()
    
    def _create_error_result(self, error_message: str, extraction_time: float) -> ExtractionResult:
        """Create error result"""
        return ExtractionResult(
            success=False,
            error_message=error_message,
            extraction_time=extraction_time
        )
    
    def _update_format_stats(self, file_format: str):
        """Update format processing statistics"""
        if file_format not in self.stats['formats_processed']:
            self.stats['formats_processed'][file_format] = 0
        self.stats['formats_processed'][file_format] += 1
    
    def _update_avg_extraction_time(self, extraction_time: float):
        """Update running average of extraction times"""
        successful_extractions = self.stats['successful_extractions']
        current_avg = self.stats['avg_extraction_time']
        
        self.stats['avg_extraction_time'] = (
            (current_avg * (successful_extractions - 1) + extraction_time) / successful_extractions
        )
    
    async def batch_extract(self, file_paths: List[str]) -> List[ExtractionResult]:
        """
        Extract text from multiple files concurrently
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            List of ExtractionResult objects
        """
        logger.info(f"📦 Starting batch text extraction: {len(file_paths)} files")
        
        tasks = [self.extract_text(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle exceptions
        final_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"❌ Batch extraction failed for {file_paths[i]}: {result}")
                final_results.append(self._create_error_result(str(result), 0))
            else:
                final_results.append(result)
        
        successful = sum(1 for r in final_results if r.success)
        logger.info(f"📦 Batch extraction completed: {successful}/{len(file_paths)} successful")
        
        return final_results
    
    def get_stats(self) -> Dict[str, Any]:
        """Get extraction statistics"""
        return {
            **self.stats,
            "ocr_enabled": self.enable_ocr,
            "ocr_language": self.ocr_language,
            "preprocessing_enabled": self.preprocessing_enabled,
            "max_image_size_mb": self.max_image_size // (1024 * 1024),
            "success_rate": (
                self.stats['successful_extractions'] / self.stats['total_extractions'] * 100 
                if self.stats['total_extractions'] > 0 else 0
            )
        }


# Global text extractor instance
text_extractor: Optional[TextExtractor] = None


def get_text_extractor() -> TextExtractor:
    """Get or create global text extractor instance"""
    global text_extractor
    
    if text_extractor is None:
        text_extractor = TextExtractor()
    
    return text_extractor